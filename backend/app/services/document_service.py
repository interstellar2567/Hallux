"""
Document Processing Service
Handles PDF, DOCX, images with OCR
Extracts text for citation verification
"""

import os
import io
from typing import Optional, List, Dict, Any
from loguru import logger
from pathlib import Path

# Optional python-magic (may not work on Windows without libmagic)
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    logger.warning("python-magic not available, will use filename extensions for file type detection")

# Document processing libraries
try:
    import PyPDF2
    from pdf2image import convert_from_bytes
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    logger.warning("PyPDF2 not available")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not available")

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    logger.warning("OCR libraries not available")


class DocumentProcessor:
    """Process various document formats and extract text"""
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'docx': ['.docx', '.doc'],
        'image': ['.jpg', '.jpeg', '.png', '.tiff', '.bmp'],
        'text': ['.txt', '.md']
    }
    
    def __init__(self):
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        
    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """
        Process uploaded document and extract text
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            enable_ocr: Whether to use OCR for images/scanned PDFs
            
        Returns:
            Dict with extracted text and metadata
        """
        logger.info(f"Processing document: {filename}")
        
        # Check file size
        if len(file_content) > self.max_file_size:
            raise ValueError(f"File too large. Max size: {self.max_file_size / 1024 / 1024}MB")
        
        # Detect file type
        file_ext = Path(filename).suffix.lower()
        
        # Extract text based on file type
        if file_ext == '.pdf':
            return await self._process_pdf(file_content, enable_ocr)
        elif file_ext in ['.docx', '.doc']:
            return await self._process_docx(file_content)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return await self._process_image(file_content)
        elif file_ext in ['.txt', '.md']:
            return await self._process_text(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    async def _process_pdf(
        self,
        file_content: bytes,
        enable_ocr: bool
    ) -> Dict[str, Any]:
        """Extract text from PDF, with OCR fallback for scanned PDFs"""
        
        if not HAS_PDF:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        
        logger.info("Extracting text from PDF...")
        
        try:
            # Try regular text extraction first
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
                    logger.debug(f"Extracted text from page {page_num + 1}/{page_count}")
            
            # If no text extracted and OCR enabled, try OCR
            if not text.strip() and enable_ocr and HAS_OCR:
                logger.info("No text found, attempting OCR...")
                text = await self._ocr_pdf(file_content)
            
            return {
                "text": text.strip(),
                "pages": page_count,
                "method": "ocr" if not text.strip() else "text_extraction",
                "success": bool(text.strip())
            }
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
    
    async def _ocr_pdf(self, file_content: bytes) -> str:
        """OCR scanned PDF pages"""
        
        if not HAS_OCR:
            raise ImportError("OCR libraries not installed. Run: pip install pytesseract pillow pdf2image")
        
        logger.info("Running OCR on PDF pages...")
        
        try:
            # Convert PDF to images
            images = convert_from_bytes(file_content)
            
            text = ""
            for i, image in enumerate(images):
                logger.debug(f"OCR on page {i + 1}/{len(images)}")
                page_text = pytesseract.image_to_string(image)
                text += page_text + "\n\n"
            
            return text
            
        except Exception as e:
            logger.error(f"OCR error: {e}")
            return ""
    
    async def _process_docx(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        logger.info("Extracting text from DOCX...")
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            # Extract paragraphs
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text])
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + "\t".join([cell.text for cell in row.cells])
            
            return {
                "text": text.strip(),
                "paragraphs": len(doc.paragraphs),
                "method": "docx_extraction",
                "success": bool(text.strip())
            }
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")
    
    async def _process_image(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        
        if not HAS_OCR:
            raise ImportError("OCR libraries not installed. Run: pip install pytesseract pillow")
        
        logger.info("Running OCR on image...")
        
        try:
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            
            return {
                "text": text.strip(),
                "dimensions": image.size,
                "method": "ocr",
                "success": bool(text.strip())
            }
            
        except Exception as e:
            logger.error(f"Image OCR error: {e}")
            raise ValueError(f"Failed to process image: {str(e)}")
    
    async def _process_text(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from plain text files"""
        
        logger.info("Reading text file...")
        
        try:
            text = file_content.decode('utf-8')
            
            return {
                "text": text.strip(),
                "method": "text_file",
                "success": bool(text.strip())
            }
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = file_content.decode(encoding)
                    return {
                        "text": text.strip(),
                        "method": "text_file",
                        "encoding": encoding,
                        "success": bool(text.strip())
                    }
                except:
                    continue
            
            raise ValueError("Failed to decode text file")


# Global document processor instance
document_processor = DocumentProcessor()
